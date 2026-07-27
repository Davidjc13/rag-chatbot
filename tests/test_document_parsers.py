"""Tests de parsers de documentos."""

from __future__ import annotations

import io

import pytest
from docx import Document
from openpyxl import Workbook

from chatbot.domain.documents import ContentKind
from chatbot.domain.exceptions import UnsupportedDocumentError
from chatbot.infrastructure.adapters.ingestion.docx_parser import DocxParserAdapter
from chatbot.infrastructure.adapters.ingestion.parser_factory import DocumentParserFactory
from chatbot.infrastructure.adapters.ingestion.xlsx_parser import XlsxParserAdapter


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Hola mundo narrativo.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Nombre"
    table.cell(0, 1).text = "Edad"
    table.cell(1, 0).text = "Ana"
    table.cell(1, 1).text = "30"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Ventas"
    sheet.append(["Producto", "Importe"])
    sheet.append(["Widget", 12.5])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def test_docx_parser_extracts_text_and_table() -> None:
    parsed = DocxParserAdapter().parse(filename="demo.docx", data=_docx_bytes())
    kinds = [b.kind for b in parsed.blocks]
    assert ContentKind.TEXT in kinds
    assert ContentKind.TABLE in kinds
    table_block = next(b for b in parsed.blocks if b.kind == ContentKind.TABLE)
    assert "| Nombre | Edad |" in table_block.text
    assert "| Ana | 30 |" in table_block.text


def test_xlsx_parser_uses_sheet_as_table() -> None:
    parsed = XlsxParserAdapter().parse(filename="demo.xlsx", data=_xlsx_bytes())
    table_block = next(b for b in parsed.blocks if b.kind == ContentKind.TABLE)
    assert "Hoja: Ventas" in table_block.text
    assert "| Producto | Importe |" in table_block.text
    assert "| Widget | 12.5 |" in table_block.text


def test_parser_factory_rejects_unknown() -> None:
    with pytest.raises(UnsupportedDocumentError):
        DocumentParserFactory().get_parser("notes.txt")


def test_parser_factory_resolves_extensions() -> None:
    factory = DocumentParserFactory()
    assert factory.get_parser("a.pdf").supports("a.pdf")
    assert factory.get_parser("a.docx").supports("a.docx")
    assert factory.get_parser("a.xlsx").supports("a.xlsx")
