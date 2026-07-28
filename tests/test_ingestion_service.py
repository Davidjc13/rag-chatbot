"""Tests del servicio de ingestión y vector store."""

from __future__ import annotations

import io

import pytest
from docx import Document

from chatbot.application.services.ingestion_service import IngestionService
from chatbot.application.services.table_aware_chunker import TableAwareChunker
from chatbot.domain.exceptions import DocumentNotFoundError, ValidationError
from chatbot.infrastructure.adapters.ingestion.parser_factory import DocumentParserFactory
from chatbot.infrastructure.adapters.llm.embedding_adapter import MockEmbeddingAdapter
from chatbot.infrastructure.adapters.persistence.memory_vector_store import InMemoryVectorStore
from chatbot.infrastructure.adapters.persistence.routed_vector_store import RoutedVectorStore


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Política de devoluciones en 30 días.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Concepto"
    table.cell(0, 1).text = "Plazo"
    table.cell(1, 0).text = "Devolución"
    table.cell(1, 1).text = "30 días"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def ingestion_service() -> IngestionService:
    return IngestionService(
        parser_factory=DocumentParserFactory(),
        chunker=TableAwareChunker(chunk_size=400, chunk_overlap=40),
        embeddings=MockEmbeddingAdapter(),
        vector_store=InMemoryVectorStore(),
    )


@pytest.mark.asyncio
async def test_ingest_docx_indexes_chunks(ingestion_service: IngestionService) -> None:
    result = await ingestion_service.ingest(filename="policy.docx", data=_docx_bytes())
    assert result.chunk_count >= 1
    assert result.filename == "policy.docx"

    docs = await ingestion_service.list_documents()
    assert len(docs) == 1
    assert docs[0].id == result.document_id


@pytest.mark.asyncio
async def test_ingest_empty_raises(ingestion_service: IngestionService) -> None:
    with pytest.raises(ValidationError):
        await ingestion_service.ingest(filename="empty.docx", data=b"")


@pytest.mark.asyncio
async def test_delete_missing_document(ingestion_service: IngestionService) -> None:
    with pytest.raises(DocumentNotFoundError):
        await ingestion_service.delete_document("missing")


@pytest.mark.asyncio
async def test_delete_document(ingestion_service: IngestionService) -> None:
    result = await ingestion_service.ingest(filename="policy.docx", data=_docx_bytes())
    await ingestion_service.delete_document(result.document_id)
    assert await ingestion_service.list_documents() == []


@pytest.mark.asyncio
async def test_ingest_mirrors_chunks_to_secondary_backend() -> None:
    primary = InMemoryVectorStore()
    secondary = InMemoryVectorStore()
    service = IngestionService(
        parser_factory=DocumentParserFactory(),
        chunker=TableAwareChunker(chunk_size=400, chunk_overlap=40),
        embeddings=MockEmbeddingAdapter(),
        vector_store=RoutedVectorStore(primary=primary, neo4j=secondary),
    )

    result = await service.ingest(filename="policy.docx", data=_docx_bytes())

    assert await primary.get_document(result.document_id) is not None
    assert await secondary.get_document(result.document_id) is not None
